
from docx import Document
from datetime import datetime
from docx.shared import Inches
import os

def generate_pdf(task_list, output_path="report.docx"):
    doc = Document()
    doc.add_heading("📋 今日任务报告", 0)

    today = datetime.now().strftime("%Y-%m-%d")
    doc.add_paragraph(f"📅 日期：{today}")

    for module, tasks in task_list.items():
        doc.add_heading(f"✅ {module} 模块任务", level=1)
        for task in tasks:
            doc.add_paragraph(f"• {task}", style="List Bullet")

    doc.save(output_path)
    print(f"✅ 任务报告已生成：{output_path}")
